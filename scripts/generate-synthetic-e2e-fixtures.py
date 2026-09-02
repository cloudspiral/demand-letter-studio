#!/usr/bin/env python3
"""Generate non-sensitive DOCX and PDF fixtures for live workflow verification."""

from __future__ import annotations

import argparse
import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def pdf_escape(value: str) -> str:
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def write_text_pdf(path: Path, title: str, lines: list[str]) -> None:
    content_lines = [
        "BT",
        "/F1 16 Tf",
        "72 742 Td",
        f"({pdf_escape(title)}) Tj",
        "/F1 11 Tf",
        "0 -28 Td",
    ]
    for line in lines:
        content_lines.extend([f"({pdf_escape(line)}) Tj", "0 -18 Td"])
    content_lines.append("ET")
    stream = "\n".join(content_lines).encode("ascii")

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    payload = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(payload))
        payload.extend(f"{index} 0 obj\n".encode("ascii"))
        payload.extend(obj)
        payload.extend(b"\nendobj\n")
    xref_offset = len(payload)
    payload.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    payload.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        payload.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    payload.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    path.write_bytes(payload)


def write_template(path: Path) -> None:
    document = Document()
    header = document.sections[0].header.paragraphs[0]
    header.text = "Claim Number: LEGACY-99999 | Mr. Obsolete Footer"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.add_heading("SYNTHETIC TIME-LIMITED POLICY LIMITS DEMAND", level=1)
    document.add_paragraph("Re: Ms. Jane Example")
    document.add_paragraph("Claim Number: LEGACY-99999")
    document.add_heading("INCIDENT AND LIABILITY", level=2)
    document.add_paragraph(
        "Ms. Jane Example was injured in a prior sample collision on January 1, 2020."
    )
    document.add_heading("MEDICAL TREATMENT", level=2)
    document.add_paragraph(
        "Ms. Jane Example received prior sample treatment and incurred $10.00 in medical expenses."
    )
    document.add_heading("DAMAGES", level=2)
    document.add_paragraph(
        "The prior sample claim included $20.00 in lost wages and requires replacement with grounded evidence."
    )
    document.add_heading("DEMAND", level=2)
    document.add_paragraph("THIS OFFER IS SUBJECT TO YOU COMPLYING WITH THE FOLLOWING EXPRESS TERMS AND CONDITIONS.")
    document.add_paragraph("SETTLEMENT CHECKS must be made payable according to counsel's written instructions.")
    document.add_paragraph("This synthetic document is for automated verification only and must never be sent.")
    document.save(path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("output_dir", type=Path)
    args = parser.parse_args()
    output_dir = args.output_dir.resolve()
    source_dir = output_dir / "sources"
    source_dir.mkdir(parents=True, exist_ok=True)

    write_template(output_dir / "synthetic-demand-template.docx")
    sources = {
        "01-incident-report.pdf": (
            "Synthetic Incident Report",
            [
                "Claimant: Jane Example",
                "Claim Number: SYNTH-2026-001",
                "Date of loss: 08/15/2026",
                "Location: 100 Test Avenue, Example City, Georgia",
                "The insured vehicle entered the intersection against a red signal.",
                "Jane Example was restrained and reported immediate neck and back pain.",
            ],
        ),
        "02-medical-record.pdf": (
            "Synthetic Medical Record",
            [
                "Patient: Jane Example",
                "Date of service: 08/15/2026",
                "Diagnosis: cervical strain and lumbar strain.",
                "Treatment: examination, radiographs, and physical therapy referral.",
                "The patient attended twelve physical therapy visits through 10/20/2026.",
                "All names and facts in this file are fictional test data.",
            ],
        ),
        "03-billing-summary.pdf": (
            "Synthetic Billing Summary",
            [
                "Patient: Jane Example",
                "Emergency department charges: $4,250.00",
                "Radiology charges: $1,800.00",
                "Physical therapy charges: $9,600.00",
                "Total medical charges: $15,650.00",
                "Balance due: $15,650.00",
            ],
        ),
        "04-insurance-declarations.pdf": (
            "Synthetic Insurance Declarations",
            [
                "Claim Number: SYNTH-2026-001",
                "Named insured: Alex Fictional",
                "Claimant: Jane Example",
                "Bodily injury liability limit: $100,000.00 per person",
                "Policy status on 08/15/2026: Active",
                "This is fictional coverage information for software verification.",
            ],
        ),
        "05-wage-statement.pdf": (
            "Synthetic Wage Statement",
            [
                "Employee: Jane Example",
                "Employer: Example Test Company",
                "Dates missed: 08/16/2026 through 08/26/2026",
                "Hours missed: 64",
                "Hourly rate: $30.00",
                "Total lost wages: $1,920.00",
            ],
        ),
    }
    for filename, (title, lines) in sources.items():
        write_text_pdf(source_dir / filename, title, lines)

    incremental_dir = output_dir / "incremental-case"
    initial_dir = incremental_dir / "initial-sources"
    initial_dir.mkdir(parents=True, exist_ok=True)
    for filename in sorted(sources)[:4]:
        shutil.copy2(source_dir / filename, initial_dir / filename)
    supplemental_path = incremental_dir / "supplemental-wage-statement.pdf"
    wage_title, wage_lines = sources["05-wage-statement.pdf"]
    write_text_pdf(supplemental_path, wage_title, wage_lines)
    (incremental_dir / "manifest.json").write_text(json.dumps({
        "scenario": "initially-incomplete-then-supplemented",
        "initialSourceCount": 4,
        "supplementalSource": supplemental_path.name,
        "expected": {
            "employee": "Jane Example",
            "hoursMissed": "64",
            "hourlyRate": "$30.00",
            "totalLostWages": "$1,920.00",
        },
        "mustNotAppearInCompleteExport": ["LEGACY-99999"],
    }, indent=2) + "\n")

    print(output_dir / "synthetic-demand-template.docx")
    print(source_dir)
    print(initial_dir)
    print(supplemental_path)


if __name__ == "__main__":
    main()
